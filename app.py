"""
Calculadora Jurídica de Execução
Escritório de Advocacia — Direito do Consumidor
Cobranças Indevidas e Fraudes Bancárias
"""

import streamlit as st
import json
import os
import re
import requests
from datetime import date, datetime
import pandas as pd
from io import BytesIO

# ──────────────────────────────────────────────────────────────
# PAGE CONFIG (deve ser a primeira chamada Streamlit)
# ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Calculadora Jurídica de Execução",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ──────────────────────────────────────────────────────────────
# CSS PERSONALIZADO
# ──────────────────────────────────────────────────────────────
st.markdown(
    """
<style>
    .main-header {
        background: linear-gradient(135deg, #1a3a6b 0%, #2c5282 100%);
        color: white;
        padding: 22px 28px;
        border-radius: 10px;
        margin-bottom: 22px;
        box-shadow: 0 4px 12px rgba(26,58,107,0.3);
    }
    .main-header h2 { margin: 0; color: white; font-size: 1.6em; }
    .main-header p  { margin: 6px 0 0 0; color: #bee3f8; font-size: 0.95em; }

    .section-header {
        background-color: #2c5282;
        color: white;
        padding: 8px 16px;
        border-radius: 6px;
        margin: 18px 0 10px 0;
        font-size: 1em;
    }

    .result-card {
        background-color: #ebf4ff;
        border: 1px solid #3182ce;
        border-radius: 8px;
        padding: 14px 18px;
        margin: 8px 0;
        font-size: 0.97em;
    }
    .result-card table { width: 100%; border-collapse: collapse; }
    .result-card tr td { padding: 3px 0; }

    .total-card {
        background: linear-gradient(135deg, #1a3a6b 0%, #2c5282 100%);
        color: white;
        border-radius: 8px;
        padding: 14px 18px;
        margin: 8px 0;
        font-size: 1.15em;
        font-weight: bold;
        box-shadow: 0 3px 10px rgba(26,58,107,0.4);
    }
    .total-card table { width: 100%; border-collapse: collapse; }

    .charge-row {
        background-color: #f7fafc;
        border: 1px solid #e2e8f0;
        border-radius: 6px;
        padding: 8px;
        margin: 4px 0;
    }

    .stButton > button {
        background-color: #2c5282;
        color: white;
        border: none;
        border-radius: 6px;
        font-weight: 500;
    }
    .stButton > button:hover {
        background-color: #1a3a6b;
        color: white;
    }
    .footer-bar {
        text-align: center;
        color: #718096;
        font-size: 0.82em;
        margin-top: 30px;
        padding-top: 15px;
        border-top: 1px solid #e2e8f0;
    }
    .col-header {
        font-weight: bold;
        font-size: 0.88em;
        color: #2c5282;
        text-align: center;
        padding: 4px 0;
        border-bottom: 2px solid #2c5282;
        margin-bottom: 6px;
    }
    .calc-value {
        font-family: monospace;
        font-size: 0.92em;
        text-align: right;
        padding: 2px 4px;
    }
</style>
""",
    unsafe_allow_html=True,
)

# Máscaras automáticas: DD/MM/AAAA para datas e 1.234,56 para valores
import streamlit.components.v1 as _components
_components.html("""
<script>
(function(){
  var setter=Object.getOwnPropertyDescriptor(window.parent.HTMLInputElement.prototype,'value').set;
  function fire(el,val){setter.call(el,val);el.dispatchEvent(new Event('input',{bubbles:true}));}

  function dateMask(input){
    if(input._dm)return; input._dm=true;
    input.addEventListener('input',function(){
      if(input._busy)return; input._busy=true;
      var d=input.value.replace(/\D/g,'').slice(0,8);
      var f=d;
      if(d.length>2)f=d.slice(0,2)+'/'+d.slice(2);
      if(d.length>4)f=d.slice(0,2)+'/'+d.slice(2,4)+'/'+d.slice(4);
      fire(input,f); input._busy=false;
    });
  }

  function currencyMask(input){
    if(input._cm)return; input._cm=true;
    input.addEventListener('input',function(){
      if(input._busy)return; input._busy=true;
      var digits=input.value.replace(/\D/g,'');
      if(!digits){fire(input,'');input._busy=false;return;}
      var num=parseInt(digits,10);
      var cents=(num%100).toString().padStart(2,'0');
      var integer=Math.floor(num/100).toString();
      integer=integer.replace(/\B(?=(\d{3})+(?!\d))/g,'.');
      fire(input,integer+','+cents);
      input._busy=false;
    });
  }

  function apply(){
    var doc=window.parent.document;
    doc.querySelectorAll('input[placeholder="DD/MM/AAAA"]').forEach(dateMask);
    doc.querySelectorAll('input[placeholder="0,00"]').forEach(currencyMask);
  }
  apply();
  new MutationObserver(apply).observe(window.parent.document.body,{childList:true,subtree:true});
})();
</script>
""", height=0)

# ──────────────────────────────────────────────────────────────
# CONSTANTES
# ──────────────────────────────────────────────────────────────
INDICES_FILE = os.path.join(os.path.dirname(__file__), "indices_juridicos.json")
BCB_API = "https://api.bcb.gov.br/dados/serie/bcdata.sgs.{serie}/dados?formato=json"

# ──────────────────────────────────────────────────────────────
# GESTÃO DE ÍNDICES (IBGE/BCB)
# ──────────────────────────────────────────────────────────────

def _parse_bcb(data: list) -> dict:
    """Converte lista BCB [{data, valor}] para dict {'YYYY-MM': float}."""
    result = {}
    for item in data:
        try:
            parts = item["data"].split("/")
            if len(parts) == 3:
                day, month, year = parts
            else:
                continue
            key = f"{year}-{month}"
            val = float(str(item["valor"]).replace(",", "."))
            result[key] = val
        except (KeyError, ValueError):
            pass
    return result


def fetch_bcb_series(serie_code: int, spinner_msg: str = "") -> dict:
    url = BCB_API.format(serie=serie_code)
    resp = requests.get(url, timeout=90)
    resp.raise_for_status()
    return _parse_bcb(resp.json())


def update_indices() -> dict:
    """Busca INPC, IPCAe e Selic do BCB e salva localmente."""
    progress = st.progress(0, text="Conectando ao Banco Central...")
    try:
        progress.progress(10, text="Buscando INPC (série 188)…")
        inpc = fetch_bcb_series(188)

        progress.progress(45, text="Buscando IPCAe (série 10764)…")
        ipcae = fetch_bcb_series(10764)

        progress.progress(75, text="Buscando Selic mensal (série 4390)…")
        selic = fetch_bcb_series(4390)

        progress.progress(95, text="Salvando localmente…")
        indices = {
            "inpc": inpc,
            "ipcae": ipcae,
            "selic": selic,
            "ultima_atualizacao": datetime.now().strftime("%d/%m/%Y %H:%M"),
        }
        with open(INDICES_FILE, "w", encoding="utf-8") as f:
            json.dump(indices, f, ensure_ascii=False, indent=2)

        progress.progress(100, text="Concluído!")
        return indices
    finally:
        progress.empty()


def load_indices() -> dict:
    if os.path.exists(INDICES_FILE):
        with open(INDICES_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"inpc": {}, "ipcae": {}, "selic": {}, "ultima_atualizacao": None}


# ──────────────────────────────────────────────────────────────
# FUNÇÕES DE CÁLCULO
# ──────────────────────────────────────────────────────────────

def month_key(year: int, month: int) -> str:
    return f"{year:04d}-{month:02d}"


def next_month(year: int, month: int):
    return (year + 1, 1) if month == 12 else (year, month + 1)


def iter_months(start: date, end: date):
    """Itera (year, month) de start até end (exclusive)."""
    cy, cm = start.year, start.month
    ey, em = end.year, end.month
    while (cy, cm) < (ey, em):
        yield cy, cm
        cy, cm = next_month(cy, cm)


def get_correction_index(year: int, month: int, indices: dict) -> float:
    """Retorna variação mensal (%) do índice de correção monetária."""
    key = month_key(year, month)
    if (year < 2024) or (year == 2024 and month <= 8):
        val = indices.get("inpc", {}).get(key)
        label = "INPC"
    else:
        val = indices.get("ipcae", {}).get(key)
        label = "IPCAe"
    return val if val is not None else 0.0


def get_interest_rate(year: int, month: int, indices: dict) -> float:
    """Retorna taxa de juros mensal (%) conforme legislação vigente."""
    if (year, month) <= (2002, 12):
        return 0.5
    elif (year, month) <= (2024, 8):
        return 1.0
    else:
        # Lei 14.905/2024 — Selic mensal (série 4390 BCB)
        val = indices.get("selic", {}).get(month_key(year, month))
        return val if val is not None else 1.0  # fallback conservador


def calculate_charge(value: float, date_charge: date, date_calc: date, indices: dict) -> dict:
    """
    Calcula correção monetária e juros simples para um lançamento.

    Período: do mês da cobrança ao mês imediatamente anterior ao cálculo.
    Juros: calculados sobre o débito corrigido (regime simples).
    """
    if date_charge >= date_calc:
        return {
            "corrected": value,
            "correction_factor": 1.0,
            "interest_pct": 0.0,
            "interest_value": 0.0,
            "total": value,
            "months": 0,
        }

    correction_factor = 1.0
    total_interest_pct = 0.0
    months_count = 0

    for year, month in iter_months(date_charge, date_calc):
        idx = get_correction_index(year, month, indices)
        correction_factor *= 1.0 + idx / 100.0
        total_interest_pct += get_interest_rate(year, month, indices)
        months_count += 1

    corrected = value * correction_factor
    interest_value = corrected * total_interest_pct / 100.0
    total = corrected + interest_value

    return {
        "corrected": corrected,
        "correction_factor": correction_factor,
        "interest_pct": total_interest_pct,
        "interest_value": interest_value,
        "total": total,
        "months": months_count,
    }


# ──────────────────────────────────────────────────────────────
# FORMATAÇÃO
# ──────────────────────────────────────────────────────────────

def fmt_brl(v: float) -> str:
    return "R$ {:,.2f}".format(v).replace(",", "X").replace(".", ",").replace("X", ".")


def fmt_pct(v: float) -> str:
    return "{:.4f}%".format(v).replace(".", ",")


def fmt_factor(v: float) -> str:
    return "{:.6f}".format(v).replace(".", ",")


# ──────────────────────────────────────────────────────────────
# EXPORT — EXCEL
# ──────────────────────────────────────────────────────────────

def export_excel(process_info: dict, rows: list, summary: dict) -> BytesIO:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    IS_INIC = process_info.get("modo") == "inicial"

    wb = Workbook()
    ws = wb.active
    ws.title = "Planilha de Débitos" if IS_INIC else "Demonstrativo"

    BLUE_DARK  = "1A3A6B"
    BLUE_MED   = "2C5282"
    BLUE_LIGHT = "EBF4FF"
    BLUE_TOTAL = "BEE3F8"
    WHITE      = "FFFFFF"

    def fill(color):
        return PatternFill("solid", fgColor=color)

    def border():
        s = Side(style="thin", color="AAAAAA")
        return Border(left=s, right=s, top=s, bottom=s)

    def font(bold=False, color="000000", size=10):
        return Font(bold=bold, color=color, size=size)

    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    right  = Alignment(horizontal="right",  vertical="center")
    left   = Alignment(horizontal="left",   vertical="center")

    row = 1

    # ── Título ──
    ws.merge_cells(f"A{row}:G{row}")
    c = ws[f"A{row}"]
    c.value = ("PLANILHA DE DÉBITOS JUDICIAIS" if IS_INIC
               else "DEMONSTRATIVO DE DÉBITO — EXECUÇÃO JUDICIAL")
    c.font = font(bold=True, color=WHITE, size=13)
    c.fill = fill(BLUE_DARK)
    c.alignment = center
    ws.row_dimensions[row].height = 24
    row += 1

    ws.merge_cells(f"A{row}:G{row}")
    c = ws[f"A{row}"]
    c.value = ("Petição Inicial — Direito do Consumidor" if IS_INIC
               else "Cumprimento de Sentença — Direito do Consumidor")
    c.font = font(color=WHITE, size=10)
    c.fill = fill(BLUE_MED)
    c.alignment = center
    row += 2

    # ── Dados do processo ──
    process_fields = [
        ("Número do Processo:", process_info.get("numero_processo", "")),
        ("Exequente:",          process_info.get("exequente", "")),
        ("Executado:",          process_info.get("executado", "")),
        ("Tribunal:",           process_info.get("tribunal", "")),
        ("Data do Cálculo:",    process_info.get("data_calculo", "")),
    ]
    for label, value in process_fields:
        ws.merge_cells(f"A{row}:B{row}")
        ws[f"A{row}"].value = label
        ws[f"A{row}"].font = font(bold=True)
        ws[f"A{row}"].fill = fill("D6E4F7")
        ws.merge_cells(f"C{row}:G{row}")
        ws[f"C{row}"].value = value
        for col in "ABCDEFG":
            if ws[f"{col}{row}"].border.left.border_style is None:
                ws[f"{col}{row}"].border = border()
        row += 1

    row += 1

    # ── Cabeçalho da tabela de lançamentos ──
    col_headers = [
        "Data da\nCobrança",
        "Valor\nOriginal (R$)",
        "Fator de\nCorreção",
        "Débito\nCorrigido (R$)",
        "Juros\nAcumulados %",
        "Valor dos\nJuros (R$)",
        "Débito\nTotal (R$)",
    ]
    col_letters = ["A", "B", "C", "D", "E", "F", "G"]
    col_widths  = [16,  18,  14,  20,  16,  18,  20]

    for i, (hdr, letter, width) in enumerate(zip(col_headers, col_letters, col_widths)):
        c = ws[f"{letter}{row}"]
        c.value = hdr
        c.font = font(bold=True, color=WHITE, size=9)
        c.fill = fill(BLUE_MED)
        c.alignment = center
        c.border = border()
        ws.column_dimensions[letter].width = width
    ws.row_dimensions[row].height = 32
    row += 1

    # ── Linhas de dados ──
    for i, r in enumerate(rows):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        values = [
            r.get("Data da Cobrança", ""),
            r.get("Valor Original", 0.0),
            r.get("Fator de Correção", 1.0),
            r.get("Débito Corrigido", 0.0),
            r.get("Total Juros %", 0.0),
            r.get("Valor dos Juros", 0.0),
            r.get("Débito Total", 0.0),
        ]
        aligns = [center, right, right, right, right, right, right]
        fmt_num = [None, '#,##0.00', '#,##0.000000', '#,##0.00', '#,##0.0000', '#,##0.00', '#,##0.00']
        for j, (letter, val, aln, nfmt) in enumerate(zip(col_letters, values, aligns, fmt_num)):
            c = ws[f"{letter}{row}"]
            c.value = val
            c.fill = fill(bg)
            c.alignment = aln
            c.border = border()
            if nfmt:
                c.number_format = nfmt
        ws.row_dimensions[row].height = 16
        row += 1

    row += 1

    # ── Resumo ──
    if IS_INIC:
        summary_rows_data = [
            ("( = ) Subtotal — Débito Corrigido:", summary["subtotal_principal"]),
            ("( = ) Subtotal — Juros Moratórios:", summary["subtotal_juros"]),
            ("( = ) Subtotal Material:",           summary["subtotal_base"]),
        ]
        if summary.get("aplicar_dobro"):
            summary_rows_data.append(("( × ) Repetição em dobro (CDC art. 42, §único):", summary["subtotal_material"]))
        if summary.get("dano_moral", 0) > 0:
            summary_rows_data.append(("( + ) Dano Moral:", summary["dano_moral"]))
        summary_rows_data.append(("VALOR DA CAUSA:", summary["total_geral"]))
        total_label_xl = "VALOR DA CAUSA:"
    else:
        summary_rows_data = [
            ("( = ) Subtotal — Débito Corrigido (Principal):", summary["subtotal_principal"]),
            ("( = ) Subtotal — Juros Moratórios:",             summary["subtotal_juros"]),
            (f"( + ) Honorários Advocatícios ({summary['honorarios_pct']:.1f}%):", summary["honorarios_valor"]),
        ]
        if summary.get("multa_523"):
            summary_rows_data.append(("( + ) Multa Art. 523 §1º CPC (10%):", summary["multa_valor"]))
        summary_rows_data.append(("( = ) TOTAL GERAL:", summary["total_geral"]))
        total_label_xl = "( = ) TOTAL GERAL:"

    for label, value in summary_rows_data:
        is_total = (label == total_label_xl)
        bg_lbl = BLUE_DARK if is_total else BLUE_TOTAL
        bg_val = BLUE_DARK if is_total else BLUE_TOTAL
        fc = WHITE if is_total else "000000"

        ws.merge_cells(f"A{row}:F{row}")
        lbl_cell = ws[f"A{row}"]
        lbl_cell.value = label
        lbl_cell.font = font(bold=True, color=fc, size=10 if not is_total else 11)
        lbl_cell.fill = fill(bg_lbl)
        lbl_cell.alignment = right
        lbl_cell.border = border()

        val_cell = ws[f"G{row}"]
        val_cell.value = value
        val_cell.number_format = '#,##0.00'
        val_cell.font = font(bold=True, color=fc, size=10 if not is_total else 11)
        val_cell.fill = fill(bg_val)
        val_cell.alignment = right
        val_cell.border = border()
        ws.row_dimensions[row].height = 18
        row += 1

    row += 2

    # ── Critérios ──
    criterios_title = ws[f"A{row}"]
    ws.merge_cells(f"A{row}:G{row}")
    criterios_title.value = "CRITÉRIOS UTILIZADOS:"
    criterios_title.font = font(bold=True, size=10)
    row += 1

    criterios = [
        "• Atualização Monetária: INPC (IBGE, série BCB 188) de jul/1995 a ago/2024; "
        "IPCAe (série BCB 10764) de set/2024 em diante.",
        "• Juros de Mora: 0,5% a.m. até dez/2002; 1% a.m. de jan/2003 a ago/2024; "
        "Taxa Selic mensal (BCB, série 4390) de set/2024 em diante (Lei 14.905/2024).",
        "• Juros calculados sobre o débito corrigido — regime de juros simples.",
        f"• Tribunal: {process_info.get('tribunal', 'TJMG')} | "
        f"Data-base do cálculo: {process_info.get('data_calculo', '')}",
    ]
    for text in criterios:
        ws.merge_cells(f"A{row}:G{row}")
        c = ws[f"A{row}"]
        c.value = text
        c.font = font(size=9)
        c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        ws.row_dimensions[row].height = 20
        row += 1

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────────────────────
# EXPORT — WORD
# ──────────────────────────────────────────────────────────────

def _word_cell_bg(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def export_word(process_info: dict, rows: list, summary: dict) -> BytesIO:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    IS_INIC = process_info.get("modo") == "inicial"
    cell_bg = _word_cell_bg

    doc = Document()

    # Fonte padrão Calibri em todo o documento
    from docx.oxml.ns import qn
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    # Margens
    sec = doc.sections[0]
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.0)

    # ── Título ──
    titulo = "PLANILHA DE DÉBITOS JUDICIAIS" if IS_INIC else "DEMONSTRATIVO DE DÉBITO"
    subtit = "PETIÇÃO INICIAL — DIREITO DO CONSUMIDOR" if IS_INIC else "CUMPRIMENTO DE SENTENÇA / EXECUÇÃO"

    t = doc.add_heading(titulo, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

    sub = doc.add_paragraph(subtit)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.color.rgb = RGBColor(0x2C, 0x52, 0x82)

    doc.add_paragraph()

    # ── Dados do processo ──
    lbl_parte = "Requerente:" if IS_INIC else "Exequente:"
    lbl_reu   = "Requerida:"  if IS_INIC else "Executado:"
    process_fields = [
        ("Número do Processo:", process_info.get("numero_processo", "")),
        (lbl_parte,             process_info.get("exequente", "")),
        (lbl_reu,               process_info.get("executado", "")),
        ("Tribunal:",           process_info.get("tribunal", "")),
        ("Data do Cálculo:",    process_info.get("data_calculo", "")),
    ]
    ptable = doc.add_table(rows=len(process_fields), cols=2)
    ptable.style = "Table Grid"
    ptable.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(process_fields):
        row = ptable.rows[i]
        cell_bg(row.cells[0], "2C5282")
        p0 = row.cells[0].paragraphs[0]
        run = p0.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(value).font.size = Pt(10)

    doc.add_paragraph()

    # ── Tabela de lançamentos ──
    tbl_headers = [
        "Data da\nCobrança",
        "Valor\nOriginal (R$)",
        "Fator de\nCorreção",
        "Débito\nCorrigido (R$)",
        "Juros\nTotal %",
        "Valor dos\nJuros (R$)",
        "Débito\nTotal (R$)",
    ]
    n_cols = len(tbl_headers)
    charges_tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    charges_tbl.style = "Table Grid"

    # Cabeçalho
    hdr_row = charges_tbl.rows[0]
    for i, hdr in enumerate(tbl_headers):
        cell = hdr_row.cells[i]
        cell_bg(cell, "1A3A6B")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8)

    # Dados
    for ri, r in enumerate(rows):
        t_row = charges_tbl.rows[ri + 1]
        bg = "EBF4FF" if ri % 2 == 0 else "FFFFFF"
        vals = [
            r.get("Data da Cobrança", ""),
            fmt_brl(r.get("Valor Original", 0.0)),
            fmt_factor(r.get("Fator de Correção", 1.0)),
            fmt_brl(r.get("Débito Corrigido", 0.0)),
            fmt_pct(r.get("Total Juros %", 0.0)),
            fmt_brl(r.get("Valor dos Juros", 0.0)),
            fmt_brl(r.get("Débito Total", 0.0)),
        ]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER] + [WD_ALIGN_PARAGRAPH.RIGHT] * (n_cols - 1)
        for ci, (val, aln) in enumerate(zip(vals, aligns)):
            cell = t_row.cells[ci]
            cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = aln
            run = p.add_run(val)
            run.font.size = Pt(8)

    doc.add_paragraph()

    # ── Resumo ──
    h = doc.add_heading("RESUMO DO CÁLCULO", 2)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

    if IS_INIC:
        sum_rows_data = [
            ("( = ) Subtotal — Débito Corrigido (Principal):", fmt_brl(summary["subtotal_principal"])),
            ("( = ) Subtotal — Juros Moratórios:",             fmt_brl(summary["subtotal_juros"])),
            ("( = ) Subtotal Material (Corrigido + Juros):",   fmt_brl(summary["subtotal_base"])),
        ]
        if summary.get("aplicar_dobro"):
            sum_rows_data.append(
                ("( × ) Repetição em dobro — CDC art. 42, §único:",
                 f"{fmt_brl(summary['subtotal_base'])} × 2 = {fmt_brl(summary['subtotal_material'])}")
            )
        if summary.get("dano_moral", 0) > 0:
            sum_rows_data.append(
                ("( + ) Dano Moral:", fmt_brl(summary["dano_moral"]))
            )
        sum_rows_data.append(("VALOR DA CAUSA:", fmt_brl(summary["total_geral"])))
        total_label = "VALOR DA CAUSA:"
    else:
        sp = summary["subtotal_principal"]
        sj = summary["subtotal_juros"]
        sb = sp + sj
        hp = summary["honorarios_pct"]
        hv = summary["honorarios_valor"]
        hp_p = sp * hp / 100.0
        hp_j = sj * hp / 100.0
        sp2 = sp + hp_p
        sj2 = sj + hp_j
        sum_rows_data = [
            ("( = ) Subtotal (principal):", f"{fmt_brl(sp)} + (juros) {fmt_brl(sj)}   {fmt_brl(sb)}"),
            (f"( + ) Honorários advocatícios {hp:.2f}%:", f"{fmt_brl(hp_p)} + - {fmt_brl(hp_j)}   {fmt_brl(hv)}"),
            ("( = ) Subtotal (principal):", f"{fmt_brl(sp2)} + (juros) {fmt_brl(sj2)}   {fmt_brl(sp2+sj2)}"),
        ]
        if summary.get("multa_523"):
            mv = summary["multa_valor"]
            mp = sp2 * 0.10
            mj = sj2 * 0.10
            sum_rows_data.append(
                ("( + ) Honorários Art. 523, §1º CPC 10%:", f"{fmt_brl(mp)} + - {fmt_brl(mj)}   {fmt_brl(mv)}")
            )
            sum_rows_data.append(
                ("( = ) Total (principal):", f"{fmt_brl(sp2+mp)} + (juros) {fmt_brl(sj2+mj)}   {fmt_brl(summary['total_geral'])}")
            )
        sum_rows_data.append(("TOTAL GERAL:", fmt_brl(summary["total_geral"])))
        total_label = "TOTAL GERAL:"

    stbl = doc.add_table(rows=len(sum_rows_data), cols=2)
    stbl.style = "Table Grid"
    stbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(sum_rows_data):
        is_total = (label == total_label)
        bg = "1A3A6B" if is_total else "BEE3F8"
        fc = RGBColor(0xFF, 0xFF, 0xFF) if is_total else RGBColor(0x00, 0x00, 0x00)
        row = stbl.rows[i]
        for cell in row.cells:
            cell_bg(cell, bg)
        p0 = row.cells[0].paragraphs[0]
        run = p0.add_run(label)
        run.bold = True
        run.font.color.rgb = fc
        run.font.size = Pt(10)
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p1.add_run(value)
        run.bold = True
        run.font.color.rgb = fc
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Critérios ──
    h2 = doc.add_heading("CRITÉRIOS UTILIZADOS", 2)
    h2.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

    criterios = [
        "Atualização Monetária: INPC (IBGE, série BCB 188) de julho/1995 a agosto/2024; "
        "IPCAe (série BCB 10764) de setembro/2024 em diante.",
        "Juros de Mora: 0,5% ao mês até dezembro/2002; 1% ao mês de janeiro/2003 a agosto/2024; "
        "Taxa Selic mensal (BCB, série 4390) de setembro/2024 em diante (Lei 14.905/2024).",
        "Juros calculados sobre o débito corrigido — regime de juros simples (não composto).",
        f"Tribunal: {process_info.get('tribunal', 'TJMG')}.",
    ]
    if IS_INIC and summary.get("aplicar_dobro"):
        criterios.append(
            "Repetição em dobro aplicada nos termos do CDC art. 42, §único "
            "(EAREsp 676.608/RS — STJ, independe de comprovação de má-fé)."
        )
    for text in criterios:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Assinatura ──
    sig1 = doc.add_paragraph(
        f"{'Local: ' + process_info.get('tribunal','') + ' — ' if not IS_INIC else 'Local e data: _________________________, '}"
        f"{process_info.get('data_calculo', '')}."
    )
    if sig1.runs:
        sig1.runs[0].font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    line = doc.add_paragraph("_" * 55)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    assinante = process_info.get("assinante_nome") or "Advogado(a) — OAB/_____ n.º _______"
    cargo     = process_info.get("assinante_cargo") or ("Subscritor do Cálculo" if IS_INIC else "Cargo / Função")
    name_line = doc.add_paragraph(f"{assinante}\n{cargo}")
    name_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in name_line.runs:
        run.font.size = Pt(10)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_pdf(process_info: dict, rows: list, summary: dict) -> BytesIO:
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT

    IS_INIC = process_info.get("modo") == "inicial"

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=(21 * cm, 29.7 * cm),
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.0 * cm,
    )

    DARK_BLUE  = colors.HexColor("#1A3A6B")
    MED_BLUE   = colors.HexColor("#2C5282")
    LIGHT_BLUE = colors.HexColor("#BEE3F8")
    ROW_EVEN   = colors.HexColor("#EBF4FF")
    WHITE      = colors.white
    BLACK      = colors.black

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "PdfTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        textColor=DARK_BLUE,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "PdfSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        textColor=MED_BLUE,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    section_style = ParagraphStyle(
        "PdfSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        textColor=DARK_BLUE,
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "PdfBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "PdfBullet",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        leftIndent=12,
        firstLineIndent=-8,
        spaceAfter=3,
    )

    story = []

    titulo = "PLANILHA DE DÉBITOS JUDICIAIS" if IS_INIC else "DEMONSTRATIVO DE DÉBITO"
    subtit = "PETIÇÃO INICIAL — DIREITO DO CONSUMIDOR" if IS_INIC else "CUMPRIMENTO DE SENTENÇA / EXECUÇÃO"
    story.append(Paragraph(titulo, title_style))
    story.append(Paragraph(subtit, subtitle_style))
    story.append(Spacer(1, 4))

    # ── Dados do processo ──
    lbl_parte = "Requerente:" if IS_INIC else "Exequente:"
    lbl_reu   = "Requerida:"  if IS_INIC else "Executado:"
    process_fields = [
        ("Número do Processo:", process_info.get("numero_processo", "")),
        (lbl_parte,             process_info.get("exequente", "")),
        (lbl_reu,               process_info.get("executado", "")),
        ("Tribunal:",           process_info.get("tribunal", "")),
        ("Data do Cálculo:",    process_info.get("data_calculo", "")),
    ]
    ptbl_data = [
        [
            Paragraph(f"<b><font color='white'>{lbl}</font></b>", body_style),
            Paragraph(val, body_style),
        ]
        for lbl, val in process_fields
    ]
    ptbl = Table(ptbl_data, colWidths=[4.5 * cm, 11.5 * cm])
    ptbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), MED_BLUE),
        ("BACKGROUND", (1, 0), (1, -1), WHITE),
        ("GRID",       (0, 0), (-1, -1), 0.3, colors.grey),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(ptbl)
    story.append(Spacer(1, 10))

    # ── Tabela de lançamentos ──
    story.append(Paragraph("LANÇAMENTOS", section_style))

    tbl_headers = [
        "Data da\nCobrança",
        "Valor\nOriginal (R$)",
        "Fator de\nCorreção",
        "Débito\nCorrigido (R$)",
        "Juros\nTotal %",
        "Valor dos\nJuros (R$)",
        "Débito\nTotal (R$)",
    ]

    hdr_cells = [
        Paragraph(f"<b><font color='white'>{h}</font></b>", ParagraphStyle(
            "ch", fontName="Helvetica-Bold", fontSize=7, alignment=TA_CENTER,
            textColor=WHITE, leading=9,
        ))
        for h in tbl_headers
    ]
    charges_data = [hdr_cells]
    for ri, r in enumerate(rows):
        vals = [
            r.get("Data da Cobrança", ""),
            fmt_brl(r.get("Valor Original", 0.0)),
            fmt_factor(r.get("Fator de Correção", 1.0)),
            fmt_brl(r.get("Débito Corrigido", 0.0)),
            fmt_pct(r.get("Total Juros %", 0.0)),
            fmt_brl(r.get("Valor dos Juros", 0.0)),
            fmt_brl(r.get("Débito Total", 0.0)),
        ]
        aligns = [TA_CENTER] + [TA_RIGHT] * 6
        charges_data.append([
            Paragraph(v, ParagraphStyle("cd", fontName="Helvetica", fontSize=7,
                                        alignment=a, leading=9))
            for v, a in zip(vals, aligns)
        ])

    col_w = [2.3 * cm, 2.1 * cm, 2.0 * cm, 2.3 * cm, 2.0 * cm, 2.1 * cm, 2.2 * cm]
    charges_tbl = Table(charges_data, colWidths=col_w, repeatRows=1)
    ts = [
        ("BACKGROUND", (0, 0), (-1, 0), DARK_BLUE),
        ("GRID",       (0, 0), (-1, -1), 0.3, colors.grey),
        ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]
    for ri in range(len(rows)):
        bg = ROW_EVEN if ri % 2 == 0 else WHITE
        ts.append(("BACKGROUND", (0, ri + 1), (-1, ri + 1), bg))
    charges_tbl.setStyle(TableStyle(ts))
    story.append(charges_tbl)
    story.append(Spacer(1, 10))

    # ── Resumo ──
    story.append(Paragraph("RESUMO DO CÁLCULO", section_style))

    if IS_INIC:
        sum_rows_data = [
            ("( = ) Subtotal — Débito Corrigido (Principal):", fmt_brl(summary["subtotal_principal"])),
            ("( = ) Subtotal — Juros Moratórios:",             fmt_brl(summary["subtotal_juros"])),
            ("( = ) Subtotal Material (Corrigido + Juros):",   fmt_brl(summary["subtotal_base"])),
        ]
        if summary.get("aplicar_dobro"):
            sum_rows_data.append((
                "( × ) Repetição em dobro — CDC art. 42, §único:",
                f"{fmt_brl(summary['subtotal_base'])} × 2 = {fmt_brl(summary['subtotal_material'])}"
            ))
        if summary.get("dano_moral", 0) > 0:
            sum_rows_data.append(("( + ) Dano Moral:", fmt_brl(summary["dano_moral"])))
        sum_rows_data.append(("VALOR DA CAUSA:", fmt_brl(summary["total_geral"])))
        total_label = "VALOR DA CAUSA:"
    else:
        sp = summary["subtotal_principal"]
        sj = summary["subtotal_juros"]
        sb = sp + sj
        hp = summary["honorarios_pct"]
        hv = summary["honorarios_valor"]
        hp_p = sp * hp / 100.0
        hp_j = sj * hp / 100.0
        sp2 = sp + hp_p
        sj2 = sj + hp_j
        sum_rows_data = [
            ("( = ) Subtotal (principal):", f"{fmt_brl(sp)} + (juros) {fmt_brl(sj)}   {fmt_brl(sb)}"),
            (f"( + ) Honorários advocatícios {hp:.2f}%:", f"{fmt_brl(hp_p)} + - {fmt_brl(hp_j)}   {fmt_brl(hv)}"),
            ("( = ) Subtotal (principal):", f"{fmt_brl(sp2)} + (juros) {fmt_brl(sj2)}   {fmt_brl(sp2+sj2)}"),
        ]
        if summary.get("multa_523"):
            mv = summary["multa_valor"]
            mp = sp2 * 0.10
            mj = sj2 * 0.10
            sum_rows_data.append((
                "( + ) Honorários Art. 523, §1º CPC 10%:", f"{fmt_brl(mp)} + - {fmt_brl(mj)}   {fmt_brl(mv)}"
            ))
            sum_rows_data.append((
                "( = ) Total (principal):", f"{fmt_brl(sp2+mp)} + (juros) {fmt_brl(sj2+mj)}   {fmt_brl(summary['total_geral'])}"
            ))
        sum_rows_data.append(("TOTAL GERAL:", fmt_brl(summary["total_geral"])))
        total_label = "TOTAL GERAL:"

    sum_pdf_data = []
    for lbl, val in sum_rows_data:
        is_total = (lbl == total_label)
        bg_color = DARK_BLUE if is_total else LIGHT_BLUE
        txt_color = "white" if is_total else "black"
        sum_pdf_data.append([
            Paragraph(f"<b><font color='{txt_color}'>{lbl}</font></b>",
                      ParagraphStyle("sl", fontName="Helvetica-Bold", fontSize=9,
                                     textColor=(WHITE if is_total else BLACK))),
            Paragraph(f"<b><font color='{txt_color}'>{val}</font></b>",
                      ParagraphStyle("sv", fontName="Helvetica-Bold", fontSize=9,
                                     alignment=TA_RIGHT,
                                     textColor=(WHITE if is_total else BLACK))),
        ])

    sum_style = []
    for si, (lbl, _) in enumerate(sum_rows_data):
        is_total = (lbl == total_label)
        bg_color = DARK_BLUE if is_total else LIGHT_BLUE
        sum_style.append(("BACKGROUND", (0, si), (-1, si), bg_color))
    sum_style += [
        ("GRID",    (0, 0), (-1, -1), 0.3, colors.grey),
        ("VALIGN",  (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    stbl = Table(sum_pdf_data, colWidths=[11 * cm, 5 * cm])
    stbl.setStyle(TableStyle(sum_style))
    story.append(stbl)
    story.append(Spacer(1, 10))

    # ── Critérios ──
    story.append(Paragraph("CRITÉRIOS UTILIZADOS", section_style))
    criterios = [
        "Atualização Monetária: INPC (IBGE, série BCB 188) de julho/1995 a agosto/2024; "
        "IPCAe (série BCB 10764) de setembro/2024 em diante.",
        "Juros de Mora: 0,5% ao mês até dezembro/2002; 1% ao mês de janeiro/2003 a agosto/2024; "
        "Taxa Selic mensal (BCB, série 4390) de setembro/2024 em diante (Lei 14.905/2024).",
        "Juros calculados sobre o débito corrigido — regime de juros simples (não composto).",
        f"Tribunal: {process_info.get('tribunal', 'TJMG')}.",
    ]
    if IS_INIC and summary.get("aplicar_dobro"):
        criterios.append(
            "Repetição em dobro aplicada nos termos do CDC art. 42, §único "
            "(EAREsp 676.608/RS — STJ, independe de comprovação de má-fé)."
        )
    for text in criterios:
        story.append(Paragraph(f"• {text}", bullet_style))

    story.append(Spacer(1, 20))

    # ── Assinatura ──
    loc_line = (
        f"Local: {process_info.get('tribunal','')} — {process_info.get('data_calculo', '')}."
        if not IS_INIC else
        f"Local e data: _________________________, {process_info.get('data_calculo', '')}."
    )
    story.append(Paragraph(loc_line, body_style))
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="60%", thickness=0.5, color=BLACK, hAlign="CENTER"))
    story.append(Spacer(1, 4))

    assinante = process_info.get("assinante_nome") or "Advogado(a) — OAB/_____ n.º _______"
    cargo     = process_info.get("assinante_cargo") or ("Subscritor do Cálculo" if IS_INIC else "Cargo / Função")
    sig_style = ParagraphStyle("sig", fontName="Helvetica", fontSize=9, alignment=TA_CENTER)
    story.append(Paragraph(assinante, sig_style))
    story.append(Paragraph(cargo, sig_style))

    doc.build(story)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────────────────────
# EXTRAÇÃO DE TEXTO E PARSING DE DOCUMENTOS
# ──────────────────────────────────────────────────────────────

_DATE_RE         = re.compile(r'\b(\d{2})[/\-\.](\d{2})[/\-\.](\d{2,4})\b')
_DATE_RE_PARTIAL = re.compile(r'\b(\d{2})[/](\d{2})\b')   # DD/MM sem ano (extratos bancários)
_VALUE_RE = re.compile(
    r'R\$\s*([\d]{1,3}(?:\.[\d]{3})*,\d{2})'          # R$ 1.234,56
    r'|(?<![,\d])([\d]{1,3}(?:\.[\d]{3})+,\d{2})(?![,\d])'  # 1.234,56 sem R$
    r'|(?<![,\d])(\d{1,6},\d{2})(?![,\d\-])'           # 999,99 sem separador (ignora - colado)
)
# Palavras que indicam que um valor é crédito/saldo e não cobrança — ignorar
_IGNORE_KEYWORDS = re.compile(
    r'\b(saldo|crédito|credito|devol|estorno|ressarc|pagament|ted\s+receb|pix\s+receb|deposito|depósito)\b',
    re.IGNORECASE,
)
# Palavras que reforçam ser cobrança
_CHARGE_KEYWORDS = re.compile(
    r'\b(tarifa|taxa|cobrança|cobranca|debito|débito|encargo|juros|multa|iof|seguro|anuidade'
    r'|mensalidade|parcela|compra|lançamento|lancamento|desconto\s+em\s+folha)\b',
    re.IGNORECASE,
)
# Meses em português para inferir ano do cabeçalho
_MONTH_NAMES_PT = {
    'janeiro':1,'fevereiro':2,'março':3,'marco':3,'abril':4,'maio':5,'junho':6,
    'julho':7,'agosto':8,'setembro':9,'outubro':10,'novembro':11,'dezembro':12,
    'jan':1,'fev':2,'mar':3,'abr':4,'mai':5,'jun':6,
    'jul':7,'ago':8,'set':9,'out':10,'nov':11,'dez':12,
}
_HEADER_MONTH_YEAR = re.compile(
    r'\b(janeiro|fevereiro|mar[çc]o|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro'
    r'|jan|fev|mar|abr|mai|jun|jul|ago|set|out|nov|dez)[/\s\-](\d{4})\b',
    re.IGNORECASE,
)
_HEADER_NUM_YEAR = re.compile(r'\b(\d{1,2})[/\-](\d{4})\b')


def _infer_year_from_text(text: str) -> int:
    """Extrai o ano do cabeçalho do documento (ex: NOVEMBRO/2018, 11/2018)."""
    for m in _HEADER_MONTH_YEAR.finditer(text):
        y = int(m.group(2))
        if 1994 <= y <= date.today().year:
            return y
    for m in _HEADER_NUM_YEAR.finditer(text):
        y = int(m.group(2))
        if 1994 <= y <= date.today().year:
            return y
    return date.today().year


def _parse_value(raw: str) -> float:
    """Converte string monetária BR para float."""
    cleaned = raw.strip().replace("R$", "").strip().replace(".", "").replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return 0.0


def _parse_date(d: str, m: str, y: str):
    try:
        year = int(y)
        if year < 100:
            year += 2000 if year < 50 else 1900
        return date(year, int(m), int(d))
    except ValueError:
        return None


def extract_text_pdf(file_bytes: bytes) -> str:
    # PyMuPDF: reconstrói linhas pela posição Y de cada palavra
    # Resolve PDFs com colunas separadas (ex: extratos bancários Itaú)
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = []
        for page in doc:
            words = page.get_text("words")  # (x0,y0,x1,y1, word, block, line, wnum)
            if not words:
                parts.append(page.get_text("text"))
                continue
            # Agrupa palavras pela posição vertical (Y) com tolerância de 4pt
            # Palavras na mesma linha visual ficam juntas, ordenadas por X (esquerda→direita)
            rows: dict[int, list] = {}
            for w in words:
                y_key = round(w[1] / 4)  # quantiza Y a grupos de 4pt
                rows.setdefault(y_key, []).append((w[0], w[4]))  # (x, texto)
            page_lines = []
            for y_key in sorted(rows):
                line_words = sorted(rows[y_key], key=lambda t: t[0])
                page_lines.append("  ".join(t[1] for t in line_words))
            parts.append("\n".join(page_lines))
        doc.close()
        text = "\n".join(parts)
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass

    # Fallback: pdfplumber
    try:
        import pdfplumber
        import io
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    except ImportError:
        return "__NO_PDFPLUMBER__"
    except Exception as e:
        return f"__ERROR__: {e}"


def extract_text_image(file_bytes: bytes) -> str:
    """OCR via EasyOCR (offline, suporta português). Faz download do modelo na 1ª vez."""
    try:
        import easyocr
        import numpy as np
        from PIL import Image
        import io

        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        img_np = np.array(img)

        # Cache do reader na session para não recarregar o modelo a cada chamada
        if "easyocr_reader" not in st.session_state:
            st.session_state.easyocr_reader = easyocr.Reader(
                ["pt", "en"],   # português + inglês
                gpu=False,
                verbose=False,
            )
        reader = st.session_state.easyocr_reader

        results = reader.readtext(img_np, detail=0, paragraph=True)
        return "\n".join(str(r) for r in results)

    except ImportError:
        return "__NO_EASYOCR__"
    except Exception as e:
        return f"__ERROR__: {e}"


def extract_text_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        return f"__ERROR__: {e}"


def parse_charges_from_text(text: str, filtro: str = "") -> list:
    """
    Extrai pares (data, valor) de texto de extratos ou sentenças.
    Se 'filtro' for fornecido, busca o termo em cada linha e também nas
    linhas vizinhas (janela de ±4 linhas) para encontrar data e valor
    mesmo quando estão em linhas separadas do nome da cobrança.
    """
    results = []
    seen = set()
    lines = [l.strip() for l in text.splitlines()]
    filtro_re = re.compile(re.escape(filtro), re.IGNORECASE) if filtro else None

    # Ano progressivo por linha: atualiza sempre que encontra um cabeçalho de mês/ano
    # Resolve extratos com múltiplas páginas e anos diferentes (ex: Itaú 190 páginas)
    _year_per_line = []
    _cur_year = date.today().year
    for _ln in lines:
        for _m in _HEADER_MONTH_YEAR.finditer(_ln):
            _y = int(_m.group(2))
            if 1994 <= _y <= date.today().year:
                _cur_year = _y
                break
        else:
            for _m in _HEADER_NUM_YEAR.finditer(_ln):
                _y = int(_m.group(2))
                if 1994 <= _y <= date.today().year:
                    _cur_year = _y
                    break
        _year_per_line.append(_cur_year)

    def _find_date_in_line(line: str, line_idx: int) -> date | None:
        """Tenta extrair data da linha: primeiro DD/MM/AAAA, depois DD/MM com ano da página."""
        for d, m, y in _DATE_RE.findall(line):
            pd_ = _parse_date(d, m, y)
            if pd_ and 1994 <= pd_.year <= date.today().year:
                return pd_
        year_ctx = _year_per_line[line_idx] if line_idx < len(_year_per_line) else date.today().year
        for d, m in _DATE_RE_PARTIAL.findall(line):
            try:
                di, mi = int(d), int(m)
                if 1 <= di <= 31 and 1 <= mi <= 12:
                    pd_ = date(year_ctx, mi, di)
                    if pd_ <= date.today():
                        return pd_
            except ValueError:
                pass
        return None

    def _find_first_value_in_line(line: str) -> float | None:
        """Retorna o primeiro valor monetário válido encontrado na linha."""
        for v_match in _VALUE_RE.findall(line):
            raw = next((v for v in v_match if v), None)
            if raw:
                val = _parse_value(raw)
                if 0 < val <= 9_999_999:
                    return val
        return None

    if filtro_re:
        # Para cada ocorrência do nome da cobrança, busca data e valor
        # mais próximos (nas linhas vizinhas se necessário)
        WINDOW = 6
        for i, line in enumerate(lines):
            if not filtro_re.search(line):
                continue

            start = max(0, i - WINDOW)
            end   = min(len(lines), i + WINDOW + 1)

            # Data mais próxima (prioriza mesma linha, depois vizinhas)
            best_date = None
            best_date_dist = WINDOW + 1
            for j in range(start, end):
                dist = abs(j - i)
                if dist >= best_date_dist:
                    continue
                pd_ = _find_date_in_line(lines[j], j)
                if pd_:
                    best_date_dist = dist
                    best_date = pd_

            # Valor mais próximo (prioriza mesma linha, depois vizinhas)
            best_val = None
            best_val_dist = WINDOW + 1
            for j in range(start, end):
                dist = abs(j - i)
                if dist >= best_val_dist:
                    continue
                val = _find_first_value_in_line(lines[j])
                if val:
                    best_val_dist = dist
                    best_val = val

            if best_date and best_val:
                key = (best_date, round(best_val, 2))
                if key not in seen:
                    seen.add(key)
                    results.append({
                        "selecionado": True,
                        "data":        best_date,
                        "valor":       best_val,
                        "descricao":   line[:120],
                        "_is_charge":  True,
                    })
    else:
        # Sem filtro: linha a linha, aceita datas completas e parciais
        for li, line in enumerate(lines):
            if not line or len(line) < 8:
                continue
            if _IGNORE_KEYWORDS.search(line):
                continue
            parsed_date = _find_date_in_line(line, li)
            if not parsed_date:
                continue
            val = _find_first_value_in_line(line)
            if not val:
                continue
            key = (parsed_date, round(val, 2))
            if key in seen:
                continue
            seen.add(key)
            is_charge = bool(_CHARGE_KEYWORDS.search(line))
            results.append({
                "selecionado": True,
                "data":        parsed_date,
                "valor":       val,
                "descricao":   line[:120],
                "_is_charge":  is_charge,
            })

    results.sort(key=lambda r: r["data"])
    return results


def parse_sentenca_from_text(text: str) -> dict:
    """
    Tenta extrair da sentença/acórdão:
    - valor da condenação principal
    - data da condenação / trânsito em julgado
    - índice de correção monetária determinado
    - taxa de juros determinada
    - honorários advocatícios (% ou valor fixo)
    - multa prevista
    Retorna dict com as chaves encontradas (valor None quando não detectado).
    """
    result = {
        "valor_condenacao": None,
        "data_condenacao": None,
        "data_transito": None,
        "indice_correcao": None,
        "taxa_juros": None,
        "honorarios_pct": None,
        "honorarios_valor": None,
        "multa_pct": None,
        "obs": [],
    }

    # Padrões de data
    _date_re = re.compile(r'\b(\d{2})[/\-\.](\d{2})[/\-\.](\d{2,4})\b')

    # Padrões de valor monetário
    _val_re = re.compile(
        r'R\$\s*([\d]{1,3}(?:\.[\d]{3})*,\d{2})'
        r'|(?<![,\d])([\d]{1,3}(?:\.[\d]{3})+,\d{2})(?![,\d])'
    )

    # Padrões de porcentagem
    _pct_re = re.compile(r'(\d{1,3}(?:[,\.]\d{1,4})?)\s*%')

    lines = text.splitlines()

    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        low = line_clean.lower()

        # ── Valor da condenação ────────────────────────────────
        if result["valor_condenacao"] is None:
            if any(kw in low for kw in [
                "condeno", "condena", "valor da condena", "valor total da condena",
                "valor do débito", "valor do debito", "importância de", "importancia de",
                "no valor de", "ao pagamento de", "a quantia de", "à quantia de",
            ]):
                for m in _val_re.finditer(line_clean):
                    raw = next((v for v in m.groups() if v), None)
                    if raw:
                        val = _parse_value(raw)
                        if 1.0 < val < 100_000_000:
                            result["valor_condenacao"] = val
                            result["obs"].append(f"Condenação: {line_clean[:120]}")
                            break

        # ── Data da condenação / publicação ───────────────────
        if result["data_condenacao"] is None:
            if any(kw in low for kw in [
                "data da senten", "prolatada em", "julgad", "publicad", "aos ", "decidid"
            ]):
                for dm in _date_re.findall(line_clean):
                    d = _parse_date(*dm)
                    if d and d.year >= 1990:
                        result["data_condenacao"] = d
                        break

        # ── Trânsito em julgado ───────────────────────────────
        if result["data_transito"] is None:
            if any(kw in low for kw in ["trânsito em julgado", "transito em julgado", "transitou em julgado"]):
                for dm in _date_re.findall(line_clean):
                    d = _parse_date(*dm)
                    if d and d.year >= 1990:
                        result["data_transito"] = d
                        break

        # ── Índice de correção ────────────────────────────────
        if result["indice_correcao"] is None:
            if any(kw in low for kw in ["inpc", "ipcae", "ipca-e", "igp-m", "igpm", "selic", "tr ", "correção monetária"]):
                for kw in ["INPC", "IPCAe", "IPCA-E", "IGP-M", "IGPM", "SELIC", "TR"]:
                    if kw.lower() in low:
                        result["indice_correcao"] = kw
                        break

        # ── Juros ─────────────────────────────────────────────
        if result["taxa_juros"] is None:
            if any(kw in low for kw in ["juros de mora", "juros moratórios", "taxa de juros"]):
                for pm in _pct_re.findall(line_clean):
                    try:
                        pct = float(pm.replace(",", "."))
                        if 0.1 <= pct <= 24:
                            result["taxa_juros"] = pct
                            break
                    except ValueError:
                        pass

        # ── Honorários ────────────────────────────────────────
        if any(kw in low for kw in ["honorários", "honorarios"]):
            # tenta pegar % primeiro
            if result["honorarios_pct"] is None:
                for pm in _pct_re.findall(line_clean):
                    try:
                        pct = float(pm.replace(",", "."))
                        if 1 <= pct <= 30:
                            result["honorarios_pct"] = pct
                            break
                    except ValueError:
                        pass
            # valor fixo
            if result["honorarios_valor"] is None:
                for m in _val_re.finditer(line_clean):
                    raw = next((v for v in m.groups() if v), None)
                    if raw:
                        val = _parse_value(raw)
                        if val > 0:
                            result["honorarios_valor"] = val
                            break

        # ── Multa ─────────────────────────────────────────────
        if result["multa_pct"] is None:
            if any(kw in low for kw in ["multa", "pena de multa", "art. 523", "art 523"]):
                for pm in _pct_re.findall(line_clean):
                    try:
                        pct = float(pm.replace(",", "."))
                        if 1 <= pct <= 100:
                            result["multa_pct"] = pct
                            break
                    except ValueError:
                        pass

    return result


# ──────────────────────────────────────────────────────────────
# SESSION STATE
# ──────────────────────────────────────────────────────────────
if "charges" not in st.session_state:
    st.session_state.charges = [{"data_cobranca": date.today(), "valor": 0.0}]

if "history" not in st.session_state:
    st.session_state.history = []

if "indices" not in st.session_state:
    loaded = load_indices()
    # Se não há índices locais (ex: Streamlit Cloud), busca automaticamente
    if not loaded.get("inpc"):
        try:
            with st.spinner("🔄 Buscando índices do Banco Central (primeira vez)…"):
                loaded = update_indices()
        except Exception:
            pass  # Continua sem índices; usuário pode tentar manualmente
    st.session_state.indices = loaded

if "upload_parsed" not in st.session_state:
    st.session_state.upload_parsed = []   # lista de dicts extraídos do último upload

if "upload_filename" not in st.session_state:
    st.session_state.upload_filename = ""

if "upload_raw_texts" not in st.session_state:
    st.session_state.upload_raw_texts = {}  # {nome_arquivo: texto} do último upload

# ──────────────────────────────────────────────────────────────
# SIDEBAR
# ──────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚖️ Calculadora Jurídica")
    st.markdown("*Escritório de Advocacia*")
    st.markdown("*Direito do Consumidor*")
    st.divider()

    indices: dict = st.session_state.indices
    ult = indices.get("ultima_atualizacao")

    if ult:
        st.success(f"**Índices carregados**\n\n📅 Atualizado em: {ult}")
        inpc_keys = sorted(indices.get("inpc", {}).keys())
        if inpc_keys:
            st.caption(f"INPC: {inpc_keys[0][:7]} → {inpc_keys[-1][:7]}")
        ipcae_keys = sorted(indices.get("ipcae", {}).keys())
        if ipcae_keys:
            st.caption(f"IPCAe: {ipcae_keys[0][:7]} → {ipcae_keys[-1][:7]}")
        selic_keys = sorted(indices.get("selic", {}).keys())
        if selic_keys:
            st.caption(f"Selic: {selic_keys[0][:7]} → {selic_keys[-1][:7]}")
    else:
        st.warning(
            "⚠️ **Índices não carregados**\n\n"
            "Clique em *Atualizar Índices* para baixar os dados do BCB/IBGE.\n\n"
            "Necessário apenas uma vez."
        )

    st.divider()

    if st.button("🔄 Atualizar Índices", use_container_width=True, type="primary"):
        try:
            new_indices = update_indices()
            st.session_state.indices = new_indices
            indices = new_indices
            st.success("✅ Índices atualizados com sucesso!")
            st.rerun()
        except Exception as e:
            st.error(f"Erro ao atualizar: {e}")

    st.divider()
    st.markdown("**📋 Histórico da Sessão**")

    if st.session_state.history:
        for item in reversed(st.session_state.history[-10:]):
            with st.container():
                st.caption(
                    f"🗂️ **{item['numero']}**\n"
                    f"{item['exequente'][:20]} vs {item['executado'][:20]}\n"
                    f"Total: {item['total']} | {item['data']}"
                )
    else:
        st.caption("Nenhum cálculo salvo ainda.")

    st.divider()
    st.caption(
        "**Legislação aplicada:**\n"
        "• Dec. 22.626/1933 (juros)\n"
        "• CC/2002, art. 406\n"
        "• Lei 14.905/2024\n"
        "• CPC/2015, art. 523"
    )


# ──────────────────────────────────────────────────────────────
# CABEÇALHO PRINCIPAL
# ──────────────────────────────────────────────────────────────
st.markdown(
    """
<div class="main-header">
  <h2>⚖️ CALCULADORA DE EXECUÇÃO JUDICIAL</h2>
  <p>Direito do Consumidor — Cobranças Indevidas e Fraudes Bancárias | TJMG / TJSP</p>
</div>
""",
    unsafe_allow_html=True,
)

indices: dict = st.session_state.indices
has_indices = bool(indices.get("inpc"))

# ── Seletor de modalidade ─────────────────────────────────────
modo_col1, modo_col2 = st.columns([2, 2])
with modo_col1:
    modo_calculo = st.radio(
        "**Modalidade do Cálculo**",
        ["📋  Petição Inicial", "⚖️  Cumprimento de Sentença / Execução"],
        horizontal=True,
        help=(
            "**Petição Inicial**: para calcular o valor da causa e anexar à petição. "
            "Inclui repetição em dobro (CDC art. 42) e dano moral.\n\n"
            "**Execução**: demonstrativo detalhado para cumprimento de sentença, "
            "com IAM, honorários e multa art. 523 CPC."
        ),
    )
IS_INICIAL   = "Inicial"    in modo_calculo
IS_EXECUCAO  = "Execução"   in modo_calculo

if not has_indices:
    st.warning(
        "⚠️ Índices monetários não carregados. Clique em **Atualizar Índices** na barra lateral "
        "para habilitar os cálculos de correção monetária (INPC/IPCAe) e juros Selic."
    )

# ──────────────────────────────────────────────────────────────
# UPLOAD DE DOCUMENTOS
# ──────────────────────────────────────────────────────────────
st.markdown(
    '<div class="section-header"><b>📎 Upload de Extrato ou Sentença — Preenchimento Automático</b></div>',
    unsafe_allow_html=True,
)

with st.expander("📂 Clique para enviar um arquivo e extrair os lançamentos automaticamente", expanded=False):
    st.markdown(
        "Faça upload de um **extrato bancário**, **sentença judicial** ou qualquer documento "
        "com datas e valores. O app irá extrair os pares _data × valor_ e pré-preencher a tabela "
        "de lançamentos. Você revisa e confirma antes de aplicar.\n\n"
        "**Formatos suportados:** PDF · PNG · JPG · DOCX\n\n"
        "> 📸 **Fotos e imagens:** lidas via **EasyOCR** (offline). "
        "Na **primeira vez** que usar imagem, o modelo de OCR (~100 MB) será baixado automaticamente — "
        "aguarde cerca de 1 minuto. Nas próximas vezes é instantâneo."
    )

    # ── 3 slots de upload lado a lado ─────────────────────────
    TIPOS = ["pdf", "png", "jpg", "jpeg", "docx"]
    slot_cols = st.columns(3)
    slot_labels = ["📄 Arquivo 1", "📄 Arquivo 2", "📄 Arquivo 3"]
    uploaded_files = []
    for col, lbl, slot_key in zip(slot_cols, slot_labels, ["up1", "up2", "up3"]):
        with col:
            st.caption(lbl)
            f = st.file_uploader(
                lbl,
                type=TIPOS,
                label_visibility="collapsed",
                key=slot_key,
            )
            uploaded_files.append(f)

    arquivos_validos = [f for f in uploaded_files if f is not None]

    st.markdown("<br>", unsafe_allow_html=True)
    filtro_col1, filtro_col2 = st.columns([2, 1])
    with filtro_col1:
        filtro_descricao = st.text_input(
            "🔎 Filtrar por nome da cobrança (opcional)",
            placeholder="Ex.: TARIFA, SEGURO, MENSALIDADE, IOF...",
            help=(
                "Se preenchido, o app extrai **apenas** as linhas do extrato que contenham esse texto. "
                "Útil para isolar uma cobrança específica quando o extrato tem vários tipos de lançamento."
            ),
        ).strip()
    with filtro_col2:
        st.markdown("<br>", unsafe_allow_html=True)
        extrair_btn = st.button(
            f"🔍  Extrair Lançamentos ({len(arquivos_validos)} arquivo(s))",
            use_container_width=True,
            type="primary",
            disabled=(len(arquivos_validos) == 0),
        )

    # ── Extração de todos os arquivos ─────────────────────────
    def _ler_arquivo(uf):
        fb  = uf.read()
        ext = uf.name.rsplit(".", 1)[-1].lower()
        if ext == "pdf":
            return extract_text_pdf(fb)
        elif ext in ("png", "jpg", "jpeg"):
            return extract_text_image(fb)
        elif ext == "docx":
            return extract_text_docx(fb)
        return ""

    if arquivos_validos and extrair_btn:
        all_parsed = []
        nomes = []
        raw_texts = {}

        for uf in arquivos_validos:
            with st.spinner(f"Lendo **{uf.name}**…"):
                raw = _ler_arquivo(uf)

            if raw == "__NO_PDFPLUMBER__":
                st.error(f"Nenhum extrator de PDF disponível. `pip install pymupdf pdfplumber`")
                raw = ""
            elif raw == "__NO_EASYOCR__":
                st.error("EasyOCR não encontrado. `pip install easyocr`")
                raw = ""
            elif raw.startswith("__ERROR__"):
                st.error(f"Erro em **{uf.name}**: {raw}")
                raw = ""

            if raw:
                parsed = parse_charges_from_text(raw, filtro=filtro_descricao)
                all_parsed.extend(parsed)
                nomes.append(uf.name)
                raw_texts[uf.name] = raw

        if all_parsed:
            # Remove duplicatas exatas (mesma data+valor de arquivos sobrepostos)
            seen = set()
            dedup = []
            for r in all_parsed:
                k = (r["data"], round(r["valor"], 2))
                if k not in seen:
                    seen.add(k)
                    dedup.append(r)
            dedup.sort(key=lambda r: r["data"])

            st.session_state.upload_parsed    = dedup
            st.session_state.upload_filename  = " + ".join(nomes)
            st.session_state.upload_raw_texts = raw_texts
            st.success(
                f"✅ **{len(dedup)} lançamentos** encontrados em "
                f"{len(nomes)} arquivo(s): *{', '.join(nomes)}*"
            )
        else:
            st.warning("Nenhum par data + valor encontrado nos arquivos enviados.")

        if raw_texts:
            with st.expander("📄 Ver texto extraído dos arquivos", expanded=True):
                st.caption("Use este texto para verificar se o app leu o PDF corretamente. "
                           "Se estiver confuso ou faltando dados, envie para análise.")
                busca_raw = st.text_input("🔍 Buscar palavra no texto extraído", placeholder="Ex: tarifa, cobrança, 01/01/2024…", key="busca_raw_text")
                for nome, txt in raw_texts.items():
                    st.markdown(f"**📎 {nome}**")
                    if busca_raw.strip():
                        linhas = txt.splitlines()
                        termo = busca_raw.strip().lower()
                        encontradas = [f"  linha {i+1}: {l}" for i, l in enumerate(linhas) if termo in l.lower()]
                        if encontradas:
                            st.success(f"{len(encontradas)} linha(s) com '{busca_raw}':")
                            st.code("\n".join(encontradas), language=None)
                        else:
                            st.warning(f"'{busca_raw}' não encontrado em {nome}.")
                    st.text_area("", txt, height=300, key=f"raw_{nome}", label_visibility="collapsed")

    # ── Preview e seleção ─────────────────────────────────────
    if st.session_state.upload_parsed:
        parsed_list = st.session_state.upload_parsed
        st.markdown(f"**Prévia — {len(parsed_list)} lançamentos extraídos de "
                    f"*{st.session_state.upload_filename}***")
        st.caption("Marque/desmarque as linhas que deseja importar e clique em **Adicionar à tabela**.")

        # Monta DataFrame editável
        preview_df = pd.DataFrame([
            {
                "✓": r["selecionado"],
                "Data":      r["data"].strftime("%d/%m/%Y"),
                "Valor (R$)": r["valor"],
                "Descrição (trecho)": r["descricao"],
            }
            for r in parsed_list
        ])

        edited = st.data_editor(
            preview_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "✓":          st.column_config.CheckboxColumn("✓",  width="small"),
                "Data":       st.column_config.TextColumn("Data",    width="small"),
                "Valor (R$)": st.column_config.NumberColumn("Valor (R$)", format="%.2f", width="medium"),
                "Descrição (trecho)": st.column_config.TextColumn("Descrição", width="large"),
            },
            key="upload_preview_editor",
        )

        imp_col1, imp_col2, imp_col3 = st.columns([1.3, 1, 2])
        with imp_col1:
            modo_import = st.radio(
                "Ao importar:",
                ["Substituir a tabela atual", "Adicionar às linhas existentes"],
                index=0,
                label_visibility="visible",
            )
        with imp_col2:
            st.markdown("<br>", unsafe_allow_html=True)
            n_sel = int(edited["✓"].sum())
            st.metric("Selecionados", f"{n_sel} de {len(edited)}")
        with imp_col3:
            st.markdown("<br>", unsafe_allow_html=True)
            aplicar = st.button(
                f"✅  Adicionar {n_sel} lançamento(s) à tabela",
                use_container_width=True,
                type="primary",
                disabled=(n_sel == 0),
            )

        if aplicar:
            novas = []
            for i, row in edited.iterrows():
                if row["✓"]:
                    try:
                        d = datetime.strptime(row["Data"], "%d/%m/%Y").date()
                    except ValueError:
                        continue
                    novas.append({"data_cobranca": d, "valor": float(row["Valor (R$)"])})

            if modo_import == "Substituir a tabela atual":
                final_charges = novas if novas else [{"data_cobranca": date.today(), "valor": 0.0}]
            else:
                existing = [c for c in st.session_state.charges if c.get("valor", 0) > 0]
                final_charges = existing + novas

            # Limpa keys antigas e pré-popula com os valores corretos
            # (evita bug do Streamlit que ignora value= quando a key já existe)
            for _k in list(st.session_state.keys()):
                if _k.startswith("date_") or _k.startswith("val_"):
                    del st.session_state[_k]
            for _i, _c in enumerate(final_charges):
                _d = _c.get("data_cobranca", date.today())
                _v = float(_c.get("valor", 0.0))
                _vs = f"{_v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                st.session_state[f"date_{_i}"] = _d.strftime("%d/%m/%Y") if isinstance(_d, date) else ""
                st.session_state[f"val_{_i}"]  = _vs
                _c["_val_str"] = _vs

            st.session_state.charges = final_charges

            st.session_state.upload_parsed   = []
            st.session_state.upload_filename = ""
            st.success(f"✅ {len(novas)} lançamentos adicionados à tabela de cálculo!")
            st.rerun()

    # ── Extração de sentença/acórdão (modo Execução) ──────────
    if IS_EXECUCAO and st.session_state.upload_raw_texts:
        sentenca_texts = st.session_state.upload_raw_texts
        all_sentenca = {}
        for nome, txt in sentenca_texts.items():
            parsed_s = parse_sentenca_from_text(txt)
            # merge: pega o primeiro valor encontrado em qualquer arquivo
            for k, v in parsed_s.items():
                if k != "obs" and v is not None and all_sentenca.get(k) is None:
                    all_sentenca[k] = v
            all_sentenca.setdefault("obs", [])
            all_sentenca["obs"].extend(parsed_s.get("obs", []))

        has_sentenca_data = any(
            all_sentenca.get(k) is not None
            for k in ["valor_condenacao", "data_condenacao", "honorarios_pct", "honorarios_valor", "taxa_juros"]
        )

        if has_sentenca_data:
            st.markdown("---")
            st.markdown("**📋 Dados extraídos da Sentença / Acórdão**")
            st.caption("Verifique os valores e use-os para preencher os campos abaixo (Resumo do Cálculo).")

            sc1, sc2, sc3, sc4 = st.columns(4)
            with sc1:
                if all_sentenca.get("valor_condenacao"):
                    st.metric("Valor da Condenação", fmt_brl(all_sentenca["valor_condenacao"]))
            with sc2:
                if all_sentenca.get("data_condenacao"):
                    st.metric("Data da Sentença", all_sentenca["data_condenacao"].strftime("%d/%m/%Y"))
                if all_sentenca.get("data_transito"):
                    st.metric("Trânsito em Julgado", all_sentenca["data_transito"].strftime("%d/%m/%Y"))
            with sc3:
                if all_sentenca.get("honorarios_pct"):
                    st.metric("Honorários (%)", f"{all_sentenca['honorarios_pct']:.0f}%")
                elif all_sentenca.get("honorarios_valor"):
                    st.metric("Honorários (R$)", fmt_brl(all_sentenca["honorarios_valor"]))
            with sc4:
                if all_sentenca.get("taxa_juros"):
                    st.metric("Juros Sentença", f"{all_sentenca['taxa_juros']:.2f}%/mês")
                if all_sentenca.get("indice_correcao"):
                    st.metric("Índice Correção", all_sentenca["indice_correcao"])
            if all_sentenca.get("multa_pct"):
                st.info(f"⚠️ Multa prevista na sentença: **{all_sentenca['multa_pct']:.0f}%**")

            if all_sentenca.get("obs"):
                with st.expander("📌 Trechos identificados na sentença", expanded=False):
                    for obs in all_sentenca["obs"]:
                        st.caption(obs)

# ──────────────────────────────────────────────────────────────
# DADOS DO PROCESSO
# ──────────────────────────────────────────────────────────────
st.markdown('<div class="section-header"><b>📄 Dados do Processo</b></div>', unsafe_allow_html=True)

col1, col2, col3 = st.columns([2, 2, 1])
with col1:
    numero_processo = st.text_input("Número do Processo", placeholder="0000000-00.0000.0.00.0000")
    exequente = st.text_input("Exequente (Credor / Autor)", placeholder="Nome completo")

with col2:
    executado = st.text_input("Executado (Devedor / Réu)", placeholder="Nome completo")
    tribunal = st.selectbox("Tribunal", ["TJMG (CGJ/MG)", "TJSP"])

with col3:
    _dc_default = date.today().strftime("%d/%m/%Y")
    dc_str = st.text_input("Data do Cálculo", value=_dc_default, placeholder="DD/MM/AAAA")
    try:
        data_calculo = datetime.strptime(dc_str.strip(), "%d/%m/%Y").date()
    except ValueError:
        data_calculo = date.today()
        if dc_str.strip():
            st.warning("⚠️ Data inválida. Use DD/MM/AAAA.")
    st.caption("Data-base para o cálculo da execução")

# ──────────────────────────────────────────────────────────────
# LANÇAMENTOS DE COBRANÇA
# ──────────────────────────────────────────────────────────────
st.markdown(
    '<div class="section-header"><b>📊 Lançamentos de Cobrança Indevida</b></div>',
    unsafe_allow_html=True,
)

# ── Criação em lote ──────────────────────────────────────────
with st.expander("⚡ Criação rápida de múltiplas linhas", expanded=False):
    st.markdown(
        "<small>Informe abaixo os valores para geração automática de lançamentos "
        "(máximo de 480 itens). O valor histórico deve ser fornecido em moeda da época, "
        "sem efetuar qualquer conversão ou corte de zeros.</small>",
        unsafe_allow_html=True,
    )
    st.markdown("")

    lc1, lc2, lc3, lc4, lc5 = st.columns([1.2, 1.5, 1.5, 2.2, 1])
    with lc1:
        lote_qtd = st.number_input(
            "Quantidade de itens", min_value=1, max_value=480, value=10, step=1,
            help="Número de linhas a criar (máx. 480)"
        )
    with lc2:
        lote_data_ini_str = st.text_input(
            "Data inicial (opcional)",
            value="",
            placeholder="DD/MM/AAAA",
            help="Deixe em branco para usar a data de hoje em todas as linhas",
        )
        try:
            lote_data_ini = datetime.strptime(lote_data_ini_str.strip(), "%d/%m/%Y").date() if lote_data_ini_str.strip() else None
        except ValueError:
            lote_data_ini = None
            if lote_data_ini_str.strip():
                st.warning("⚠️ Data inválida. Use DD/MM/AAAA.")
    with lc3:
        lote_valor = st.number_input(
            "Valor padrão (R$)", min_value=0.0, value=0.0, step=0.01, format="%.2f",
            help="Valor aplicado a todas as linhas (pode editar depois)"
        )
    with lc4:
        lote_freq = st.radio(
            "Periodicidade",
            options=["Diário", "Semanal", "Mensal", "Anual"],
            index=2,
            horizontal=True,
            help="Intervalo entre as datas geradas automaticamente",
        )
    with lc5:
        st.markdown("<div style='padding-top:28px'>", unsafe_allow_html=True)
        gerar = st.button("✅  Ok — Gerar", use_container_width=True, type="primary")
        st.markdown("</div>", unsafe_allow_html=True)

    if gerar:
        from datetime import timedelta

        base = lote_data_ini if lote_data_ini else date.today()
        novas = []
        for k in range(int(lote_qtd)):
            if lote_freq == "Diário":
                d_item = base + timedelta(days=k)
            elif lote_freq == "Semanal":
                d_item = base + timedelta(weeks=k)
            elif lote_freq == "Mensal":
                # avança k meses sem ultrapassar fim do mês
                m_total = base.month - 1 + k
                ano_k   = base.year + m_total // 12
                mes_k   = m_total % 12 + 1
                import calendar
                ultimo_dia = calendar.monthrange(ano_k, mes_k)[1]
                d_item = date(ano_k, mes_k, min(base.day, ultimo_dia))
            else:  # Anual
                try:
                    d_item = date(base.year + k, base.month, base.day)
                except ValueError:
                    d_item = date(base.year + k, base.month, 28)
            novas.append({"data_cobranca": d_item, "valor": float(lote_valor)})

        for _k in list(st.session_state.keys()):
            if _k.startswith("date_") or _k.startswith("val_"):
                del st.session_state[_k]
        for _i, _c in enumerate(novas):
            _d = _c.get("data_cobranca", date.today())
            _v = float(_c.get("valor", 0.0))
            _vs = f"{_v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            st.session_state[f"date_{_i}"] = _d.strftime("%d/%m/%Y") if isinstance(_d, date) else ""
            st.session_state[f"val_{_i}"]  = _vs
            _c["_val_str"] = _vs
        st.session_state.charges = novas
        st.success(f"✅ {int(lote_qtd)} linhas criadas com periodicidade **{lote_freq}** "
                   f"a partir de **{base.strftime('%d/%m/%Y')}**.")
        st.rerun()

# Botões de controle
btn_col1, btn_col2, btn_col3 = st.columns([1, 1, 4])
with btn_col1:
    if st.button("➕  Adicionar Linha", use_container_width=True):
        st.session_state.charges.append({"data_cobranca": date.today(), "valor": 0.0})
        st.rerun()
with btn_col2:
    if st.button("🗑️  Limpar Tudo", use_container_width=True):
        st.session_state.charges = [{"data_cobranca": date.today(), "valor": 0.0}]
        st.rerun()

# Cabeçalhos das colunas
HDR_COLS = [0.4, 2.0, 1.5, 1.6, 1.4, 1.6, 1.6, 0.5]
hcols = st.columns(HDR_COLS)
labels = ["#", "Data da Cobrança", "Valor Original (R$)",
          "Débito Corrigido", "Juros Acum. %", "Valor dos Juros", "Débito Total", "Del."]
for hc, lbl in zip(hcols, labels):
    hc.markdown(f'<div class="col-header">{lbl}</div>', unsafe_allow_html=True)

st.markdown("<hr style='margin:4px 0; border-color:#e2e8f0;'>", unsafe_allow_html=True)

# Processar cada lançamento
charges_results = []
to_remove = []

for i, charge in enumerate(st.session_state.charges):
    cols = st.columns(HDR_COLS)

    with cols[0]:
        st.markdown(f"<div style='text-align:center;padding-top:32px;color:#2c5282;font-weight:bold'>{i+1}</div>",
                    unsafe_allow_html=True)

    with cols[1]:
        _d_obj = charge.get("data_cobranca", date.today())
        _d_default_str = _d_obj.strftime("%d/%m/%Y") if isinstance(_d_obj, date) else date.today().strftime("%d/%m/%Y")
        d_str = st.text_input(
            f"data_{i}",
            value=charge.get("_date_str", _d_default_str),
            placeholder="DD/MM/AAAA",
            label_visibility="collapsed",
            key=f"date_{i}",
        )
        try:
            d = datetime.strptime(d_str.strip(), "%d/%m/%Y").date()
        except ValueError:
            d = _d_obj if isinstance(_d_obj, date) else date.today()
        st.session_state.charges[i]["data_cobranca"] = d
        st.session_state.charges[i]["_date_str"] = d_str

    with cols[2]:
        _v_raw = charge.get("valor", 0.0)
        _v_default_str = f"{float(_v_raw):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        v_str = st.text_input(
            f"val_{i}",
            value=charge.get("_val_str", _v_default_str),
            placeholder="0,00",
            label_visibility="collapsed",
            key=f"val_{i}",
        )
        try:
            v = float(v_str.strip().replace(".", "").replace(",", "."))
            if v < 0:
                v = 0.0
        except ValueError:
            v = 0.0
        st.session_state.charges[i]["valor"] = v
        st.session_state.charges[i]["_val_str"] = v_str

    # Cálculo
    if v > 0 and has_indices:
        res = calculate_charge(v, d, data_calculo, indices)
    elif v > 0:
        res = {
            "corrected": v, "correction_factor": 1.0,
            "interest_pct": 0.0, "interest_value": 0.0,
            "total": v, "months": 0,
        }
    else:
        res = {
            "corrected": 0.0, "correction_factor": 1.0,
            "interest_pct": 0.0, "interest_value": 0.0,
            "total": 0.0, "months": 0,
        }

    charges_results.append({
        "idx": i,
        "data_cobranca": d,
        "valor_original": v,
        **res,
    })

    with cols[3]:
        st.markdown(f'<div class="calc-value" style="padding-top:32px">{fmt_brl(res["corrected"])}</div>',
                    unsafe_allow_html=True)

    with cols[4]:
        st.markdown(f'<div class="calc-value" style="padding-top:32px">{fmt_pct(res["interest_pct"])}</div>',
                    unsafe_allow_html=True)

    with cols[5]:
        st.markdown(f'<div class="calc-value" style="padding-top:32px">{fmt_brl(res["interest_value"])}</div>',
                    unsafe_allow_html=True)

    with cols[6]:
        st.markdown(f'<div class="calc-value" style="padding-top:32px;font-weight:bold;color:#1a3a6b">'
                    f'{fmt_brl(res["total"])}</div>',
                    unsafe_allow_html=True)

    with cols[7]:
        st.markdown("<div style='padding-top:24px'>", unsafe_allow_html=True)
        if st.button("✕", key=f"del_{i}", help="Remover linha") and len(st.session_state.charges) > 1:
            to_remove.append(i)
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<hr style='margin:2px 0; border-color:#f0f4f8;'>", unsafe_allow_html=True)

if to_remove:
    st.session_state.charges = [
        c for j, c in enumerate(st.session_state.charges) if j not in to_remove
    ]
    st.rerun()

# ──────────────────────────────────────────────────────────────
# RESUMO E TOTAIS  (modo-dependente)
# ──────────────────────────────────────────────────────────────
st.markdown('<div class="section-header"><b>📈 Resumo do Cálculo</b></div>', unsafe_allow_html=True)

subtotal_principal = sum(r["corrected"]      for r in charges_results)
subtotal_juros     = sum(r["interest_value"] for r in charges_results)
subtotal_base      = subtotal_principal + subtotal_juros

col_ctrl, col_result = st.columns([1.3, 1])

# ── Controles específicos por modo ───────────────────────────
if IS_INICIAL:
    with col_ctrl:
        st.markdown("**🔁 Repetição de Indébito**")
        aplicar_dobro = st.checkbox(
            "Repetição em dobro — CDC art. 42, §único",
            value=True,
            help="Aplica o fator ×2 sobre o total corrigido (cobranças indevidas com má-fé do fornecedor).",
        )
        fator_dobro = 2.0 if aplicar_dobro else 1.0

        st.markdown("**⚖️ Dano Moral**")
        dano_moral = st.number_input(
            "Valor do dano moral (R$)",
            min_value=0.0, value=0.0, step=100.0, format="%.2f",
            help="Valor fixo pleiteado a título de danos morais (não sofre correção na petição inicial).",
        )
        st.markdown("---")
        st.caption(
            "**Valor da causa** = (débito corrigido + juros) × fator dobro + dano moral"
        )

    subtotal_material   = subtotal_base * fator_dobro
    total_geral         = subtotal_material + dano_moral
    honorarios_pct      = 0.0
    honorarios_valor    = 0.0
    aplicar_multa523    = False
    multa_valor         = 0.0
    assinante_nome      = ""
    assinante_cargo     = ""

    with col_result:
        dobro_row = (
            f"<tr><td>Fator repetição em dobro (×2):</td>"
            f"<td align='right'><b>{fmt_brl(subtotal_base)} × 2</b></td></tr>"
            if aplicar_dobro else ""
        )
        dano_row = (
            f"<tr><td>Dano moral:</td>"
            f"<td align='right'><b>{fmt_brl(dano_moral)}</b></td></tr>"
            if dano_moral > 0 else ""
        )
        st.markdown(
            f"""
<div class="result-card">
  <table>
    <tr><td>Subtotal Principal (Corrigido):</td>
        <td align='right'><b>{fmt_brl(subtotal_principal)}</b></td></tr>
    <tr><td>Subtotal Juros Moratórios:</td>
        <td align='right'><b>{fmt_brl(subtotal_juros)}</b></td></tr>
    <tr style='border-top:1px solid #3182ce'>
        <td><b>Subtotal Material (Corr. + Juros):</b></td>
        <td align='right'><b>{fmt_brl(subtotal_base)}</b></td></tr>
    {dobro_row}
    <tr><td><b>Total Material{'  (em dobro)' if aplicar_dobro else ''}:</b></td>
        <td align='right'><b>{fmt_brl(subtotal_material)}</b></td></tr>
    {dano_row}
  </table>
</div>
<div class="total-card">
  <table>
    <tr><td>VALOR DA CAUSA:</td>
        <td align='right'>{fmt_brl(total_geral)}</td></tr>
  </table>
</div>
""",
            unsafe_allow_html=True,
        )

else:  # IS_EXECUCAO
    with col_ctrl:
        honorarios_pct = st.number_input(
            "Honorários Advocatícios (%)",
            min_value=0.0, max_value=100.0,
            value=10.0, step=0.5, format="%.1f",
        )
        aplicar_multa523 = st.checkbox(
            "Multa Art. 523 §1º CPC (10%)",
            value=False,
            help="Multa de 10% por inadimplemento no prazo de 15 dias após intimação.",
        )
        st.markdown("---")
        st.markdown("**Assinante do Cálculo**")
        assinante_nome  = st.text_input("Nome", placeholder="Ex.: Maria da Silva")
        assinante_cargo = st.text_input("Cargo", placeholder="Ex.: Advogada — OAB/MG n.º 12345")

    aplicar_dobro    = False
    fator_dobro      = 1.0
    dano_moral       = 0.0
    subtotal_material = subtotal_base

    # Subtotais TJMG separados: principal × juros
    honor_princ  = subtotal_principal * honorarios_pct / 100.0
    honor_juros  = subtotal_juros     * honorarios_pct / 100.0
    honorarios_valor = honor_princ + honor_juros

    sub2_princ = subtotal_principal + honor_princ
    sub2_juros = subtotal_juros     + honor_juros
    sub2_total = sub2_princ + sub2_juros

    multa_princ  = sub2_princ * 0.10 if aplicar_multa523 else 0.0
    multa_juros  = sub2_juros * 0.10 if aplicar_multa523 else 0.0
    multa_valor  = multa_princ + multa_juros

    total_princ  = sub2_princ + multa_princ
    total_juros  = sub2_juros + multa_juros
    total_geral  = total_princ + total_juros

    with col_result:
        honor_row = (
            f"<tr><td>( + ) Honorários ({honorarios_pct:.1f}%):</td>"
            f"<td align='right'>princ. <b>{fmt_brl(honor_princ)}</b>"
            f" + juros <b>{fmt_brl(honor_juros)}</b>"
            f" = <b>{fmt_brl(honorarios_valor)}</b></td></tr>"
        )
        multa_row = (
            f"<tr><td>( + ) Multa Art. 523 §1º CPC (10%):</td>"
            f"<td align='right'>princ. <b>{fmt_brl(multa_princ)}</b>"
            f" + juros <b>{fmt_brl(multa_juros)}</b>"
            f" = <b>{fmt_brl(multa_valor)}</b></td></tr>"
            if aplicar_multa523 else ""
        )
        st.markdown(
            f"""
<div class="result-card">
  <table>
    <tr><td>( = ) Subtotal:</td>
        <td align='right'>princ. <b>{fmt_brl(subtotal_principal)}</b>
        + juros <b>{fmt_brl(subtotal_juros)}</b>
        = <b>{fmt_brl(subtotal_base)}</b></td></tr>
    {honor_row}
    <tr style='border-top:1px solid #3182ce'>
        <td>( = ) Subtotal c/ honorários:</td>
        <td align='right'>princ. <b>{fmt_brl(sub2_princ)}</b>
        + juros <b>{fmt_brl(sub2_juros)}</b>
        = <b>{fmt_brl(sub2_total)}</b></td></tr>
    {multa_row}
  </table>
</div>
<div class="total-card">
  <table>
    <tr><td>( = ) TOTAL GERAL:</td>
        <td align='right'>princ. {fmt_brl(total_princ)}
        + juros {fmt_brl(total_juros)}
        = {fmt_brl(total_geral)}</td></tr>
  </table>
</div>
""",
            unsafe_allow_html=True,
        )

# ──────────────────────────────────────────────────────────────
# CRITÉRIOS E TABELA DETALHADA
# ──────────────────────────────────────────────────────────────
with st.expander("📚 Critérios Aplicados no Cálculo", expanded=False):
    base_txt = (
        "**Correção Monetária:** INPC (BCB série 188) jul/1995–ago/2024; "
        "IPCAe (BCB série 10764) set/2024 em diante.\n\n"
        "**Juros de Mora (simples):** 0,5% a.m. até dez/2002; 1% a.m. jan/2003–ago/2024; "
        "Selic mensal (BCB série 4390) set/2024 em diante — Lei 14.905/2024.\n\n"
        "**Método:** IAM = produto das variações mensais INPC/IPCAe. "
        "Juros = soma simples das taxas × débito corrigido."
    )
    if IS_INICIAL:
        extra = (
            "\n\n**Repetição de Indébito:** fator × 2 aplicado sobre "
            "o total corrigido (CDC art. 42, §único — independe de comprovação de má-fé: EAREsp 676.608/RS)."
        )
    else:
        extra = (
            f"\n\n**Honorários:** {honorarios_pct:.1f}% sobre o débito corrigido + juros.\n\n"
            "**Multa Art. 523 §1º CPC:** 10% por inadimplemento após intimação para pagamento "
            "voluntário em 15 dias."
        )
    st.markdown(base_txt + extra)

with st.expander("🔎 Ver Tabela Detalhada dos Lançamentos", expanded=False):
    if not any(r["valor_original"] > 0 for r in charges_results):
        st.info("Nenhum lançamento com valor maior que zero.")
    else:
        display_data = []
        for r in charges_results:
            row_d = {
                "Data":            r["data_cobranca"].strftime("%d/%m/%Y"),
                "Valor Original":  fmt_brl(r["valor_original"]),
                "IAM":             fmt_factor(r["correction_factor"]),
                "Déb. Corrigido":  fmt_brl(r["corrected"]),
                "Juros %":         fmt_pct(r["interest_pct"]),
                "Vl. Juros":       fmt_brl(r["interest_value"]),
                "Déb. Total":      fmt_brl(r["total"]),
            }
            if IS_INICIAL and aplicar_dobro:
                row_d["Em Dobro"] = fmt_brl(r["total"] * 2)
            display_data.append(row_d)
        st.dataframe(pd.DataFrame(display_data), use_container_width=True, hide_index=True)

# ──────────────────────────────────────────────────────────────
# EXPORTAÇÃO
# ──────────────────────────────────────────────────────────────
st.markdown('<div class="section-header"><b>📤 Exportar Demonstrativo</b></div>', unsafe_allow_html=True)

process_info_export = {
    "numero_processo": numero_processo or "(não informado)",
    "exequente":       exequente or "(não informado)",
    "executado":       executado or "(não informado)",
    "tribunal":        tribunal,
    "data_calculo":    data_calculo.strftime("%d/%m/%Y"),
    "modo":            "inicial" if IS_INICIAL else "execucao",
    "assinante_nome":  assinante_nome  if IS_EXECUCAO else "",
    "assinante_cargo": assinante_cargo if IS_EXECUCAO else "",
}
summary_export = {
    "subtotal_principal":  subtotal_principal,
    "subtotal_juros":      subtotal_juros,
    "subtotal_base":       subtotal_base,
    # Inicial
    "aplicar_dobro":       aplicar_dobro,
    "fator_dobro":         fator_dobro,
    "subtotal_material":   subtotal_material,
    "dano_moral":          dano_moral,
    # Execução
    "honorarios_pct":      honorarios_pct,
    "honorarios_valor":    honorarios_valor,
    "multa_523":           aplicar_multa523,
    "multa_valor":         multa_valor,
    "total_geral":         total_geral,
}
export_rows = [
    {
        "Data da Cobrança":  r["data_cobranca"].strftime("%d/%m/%Y"),
        "Valor Original":    r["valor_original"],
        "Fator de Correção": r["correction_factor"],
        "Débito Corrigido":  r["corrected"],
        "Total Juros %":     r["interest_pct"],
        "Valor dos Juros":   r["interest_value"],
        "Débito Total":      r["total"],
    }
    for r in charges_results
]

safe_num = (numero_processo or "processo").replace("/", "_").replace(".", "_").replace(" ", "_")
date_str  = data_calculo.strftime("%Y%m%d")
modo_str  = "inicial" if IS_INICIAL else "execucao"

exp_col1, exp_col2, exp_col3, exp_col4 = st.columns(4)

with exp_col1:
    try:
        from docx import Document as _Doc  # noqa
        word_buf = export_word(process_info_export, export_rows, summary_export)
        st.download_button(
            label="📄  Exportar Word (.docx)",
            data=word_buf,
            file_name=f"calculo_{modo_str}_{safe_num}_{date_str}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    except ImportError:
        st.error("Instale: `pip install python-docx`")
    except Exception as e:
        st.error(f"Erro Word: {e}")

with exp_col2:
    try:
        import openpyxl as _xl  # noqa
        excel_buf = export_excel(process_info_export, export_rows, summary_export)
        st.download_button(
            label="📊  Exportar Excel (.xlsx)",
            data=excel_buf,
            file_name=f"calculo_{modo_str}_{safe_num}_{date_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    except ImportError:
        st.error("Instale: `pip install openpyxl`")
    except Exception as e:
        st.error(f"Erro Excel: {e}")

with exp_col3:
    try:
        from reportlab.platypus import SimpleDocTemplate  # noqa
        pdf_buf = export_pdf(process_info_export, export_rows, summary_export)
        st.download_button(
            label="🖨️  Exportar PDF (.pdf)",
            data=pdf_buf,
            file_name=f"calculo_{modo_str}_{safe_num}_{date_str}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    except ImportError:
        st.error("Instale: `pip install reportlab`")
    except Exception as e:
        st.error(f"Erro PDF: {e}")

with exp_col4:
    if st.button("💾  Salvar no Histórico", use_container_width=True):
        if numero_processo:
            st.session_state.history.append({
                "numero":    numero_processo,
                "modo":      "Inicial" if IS_INICIAL else "Execução",
                "exequente": exequente or "—",
                "executado": executado or "—",
                "data":      data_calculo.strftime("%d/%m/%Y"),
                "total":     fmt_brl(total_geral),
            })
            st.success("✅ Salvo no histórico!")
            st.rerun()
        else:
            st.warning("Informe o número do processo antes de salvar.")

# ──────────────────────────────────────────────────────────────
# RODAPÉ
# ──────────────────────────────────────────────────────────────
st.markdown(
    """
<div class="footer-bar">
  <p>
    Calculadora Jurídica v2.0 &nbsp;|&nbsp; Uso Interno — Escritório de Advocacia<br>
    Petição Inicial · Cumprimento de Sentença / Execução | TJMG / TJSP<br>
    INPC / IPCAe / Selic (BCB) | Lei 14.905/2024 | CDC art. 42<br>
    <em>Os valores calculados são de uso interno e devem ser conferidos pelo profissional responsável.</em>
  </p>
</div>
""",
    unsafe_allow_html=True,
)
